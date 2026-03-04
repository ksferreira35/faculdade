import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from io import BytesIO




def carregar_html(url):
   """
   Carrega o conteúdo HTML de uma URL.  Args: url: A URL da página HTML.
   Returns:  O conteúdo HTML da página como uma string, ou None em caso de erro.
   """
   try:
       response = requests.get(url, verify=False)
       response.raise_for_status()
       return response.text
   except requests.exceptions.RequestException as e:
       print(f"Erro ao acessar a URL {url}: {e}")
       return None




def extrair_urls_imagens(html_content):
   """
   Extrai URLs de imagens de um conteúdo HTML. Args: html_content: O conteúdo HTML como uma string.
   Returns: Uma lista de URLs de imagens encontradas no HTML.
   """
   if not html_content:
       return []


   soup = BeautifulSoup(html_content, 'html.parser')
   img_tags = soup.find_all('img')
   img_urls = [img.get('src') for img in img_tags if img.get('src')]
   return img_urls




def gerar_docx(urls_imagens, url_base):
   """
   Gera um documento .docx com uma tabela contendo as imagens, os links das imagens e a URL base.
   Args: urls_imagens: Uma lista de URLs de imagens.
         url_base: A URL base da página HTML.
   """
   documento = Document()
   tabela = documento.add_table(rows=1, cols=2)
   tabela.style = 'Table Grid'  # Adiciona bordas à tabela
   tabela.columns[0].width = Inches(3)
   tabela.columns[0].width = Inches(5)


   # Cabeçalho da tabela
   cabecalho_celulas = tabela.rows[0].cells
   cabecalho_celulas[0].text = "Imagem"
   cabecalho_celulas[1].text = "Link"


   for url_imagem in urls_imagens:
       try:
           # Baixa a imagem
           if url_imagem.__contains__('http'):
               endereco = url_imagem
           else:
               endereco = url_base + "/" + url_imagem
           response = requests.get(endereco, stream=True, verify=False)
           response.raise_for_status()


           # Adiciona uma nova linha à tabela
           linha_celulas = tabela.add_row().cells


           # Adiciona a imagem à primeira célula
           linha_celulas[0].paragraphs[0].text = ""  # Limpa o parágrafo existente
           linha_celulas[0].paragraphs[0].alignment = 1  # Centraliza a imagem
           linha_celulas[0].paragraphs[0].add_run().add_picture(BytesIO(response.content), width=Inches(3))


           # Adiciona o link da imagem à segunda célula
           linha_celulas[1].text = endereco
       except requests.exceptions.RequestException as e:
           print(f"Erro ao processar a imagem {url_imagem}: {e}")
       except Exception as e:
           print(f"Ocorreu um erro ao processar a imagem {url_imagem}: {e}")


   documento.save("imagens_e_links.docx")
   print("Documento 'imagens_e_links.docx' criado com sucesso.")




if __name__ == "__main__":
   # Bloco principal
   # https://alelo.cenargen.ambrapa.br
   # https://www.df.gov.br
   # https://www.seduh.df.gov.br
   # https://www.agenciabrasilia.df.gov.br
   url_html = "https://www.agenciabrasilia.df.gov.br"  # Substitua pela URL da página que você deseja carregar
   html_content = carregar_html(url_html)


   if html_content:
       print(html_content)
       links = extrair_urls_imagens(html_content)
       if len(links) > 0:
           gerar_docx(links, url_html)
   else:
       # Ocorreu um erro ao carregar o HTML
       print("Não foi possível carregar o conteúdo HTML.")


